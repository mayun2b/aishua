from __future__ import annotations

from datetime import datetime
import math
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:  # pragma: no cover - fallback for environments without Pillow
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT_DIR = Path(__file__).resolve().parents[1]
DOCS_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = DOCS_DIR / "基于SpringBoot+Vue3的AI刷题系统设计与实现-项目模板.docx"
TEMPLATE_DIR = Path(r"E:\download")
GENERATED_FIGURE_DIR = DOCS_DIR / "generated_figures"
CAPTURED_FIGURE_DIR = DOCS_DIR / "captured_figures"
TITLE = "基于SpringBoot+Vue3的AI刷题系统设计与实现"
REGULAR_FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\simfang.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
    Path(r"C:\Windows\Fonts\simsunb.ttf"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
]
BOLD_FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsunb.ttf"),
    Path(r"C:\Windows\Fonts\simkai.ttf"),
]

ABSTRACT = (
    "本文围绕基于SpringBoot与Vue3的AI刷题系统展开设计与实现研究。该系统以学科练习、错题复盘、模拟考试和学情分析为核心场景，"
    "旨在解决传统刷题平台在题目组织松散、练习结果沉淀不足、错题反馈浅层化以及智能分析能力脱离主业务链路等方面的不足。"
    "系统后端采用SpringBoot 3.1.8构建，结合MyBatis-Plus实现数据访问，使用MySQL保存题库、练习、考试和AI分析等核心业务数据，"
    "通过Redis完成练习草稿缓存与短时状态存储，借助MinIO管理题图、主观题作答图片等对象文件资源；前端采用Vue3、Vue Router、Vuex、Vant与ECharts实现学生端和后台管理端页面。"
    "在智能能力方面，系统对接千问兼容接口与Dify工作流，实现错题AI分析、练习题智能助教、薄弱知识点视频检索和学习分析报告等功能。"
    "本文围绕系统需求分析、总体设计、数据库设计、核心模块实现以及部署验证方案进行了系统阐述。"
    "实践结果表明，该系统已经形成“加入学科、开始练习、提交判题、沉淀错题、智能分析、持续强化”的学习闭环，"
    "具有较好的课程实践价值和进一步扩展为智能学习助手平台的基础。"
)
KEYWORDS = "关键词：SpringBoot，Vue3，AI刷题系统，错题分析，Dify"

TOC_REPLACEMENTS = {
    "2.2.2 uni-app框架概述\t5": "2.2.2 Vue3框架概述\t4",
    "2.2.5 RabbitMQ概述\t6": "2.2.5 MinIO对象存储概述\t5",
    "2.2.6 大模型接口概述\t6": "2.2.6 Dify与大模型接口概述\t6",
    "3.3 非功能需求\t20": "3.4 非功能需求\t9",
    "3.3.1 环境需求\t20": "3.4.1 环境需求\t9",
    "3.3.2 安全性需求\t20": "3.4.2 安全性需求\t9",
    "3.3.3 用户体验需求\t20": "3.4.3 用户体验需求\t10",
    "3.3.4 性能需求\t21": "3.4.4 性能需求\t10",
    "4.2.1 用户管理模块\t23": "4.2.1 用户与权限模块\t13",
    "4.2.2 志愿者管理模块\t23": "4.2.2 学科与目录模块\t13",
    "4.2.3 视障求助通话模块\t23": "4.2.3 练习闭环模块\t14",
    "4.2.4 政策管理模块\t23": "4.2.4 AI智能分析模块\t14",
    "4.2.6 无障碍适配模块\t23": "4.2.5 模拟考试模块\t15",
    "5.1.1 首页功能实现\t26": "5.1.1 学生端学科与首页功能实现\t18",
    "5.1.2 视障求助通话功能实现\t27": "5.1.2 学生端练习、错题与AI助教功能实现\t19",
    "5.1.3 政策查询功能实现\t27": "5.1.3 学生端模拟考试功能实现\t20",
    "5.1.4 志愿者服务模块实现\t27": "5.1.4 学生端学情分析与薄弱点推荐实现\t21",
    "5.1.5 后台管理模块实现\t27": "5.1.5 管理员端学科、目录与考点管理实现\t22",
    "5.1.6 无障碍适配功能实现\t27": "5.1.6 管理员端题库、试卷与考试记录管理实现\t23",
    "5.1.7 智能助手功能实现\t27": "5.1.7 文件上传与资源访问实现\t24",
    "6.3.3 部署验证\t28": "6.3 部署验证\t24",
}

SCHEMA_SQL_DIR = ROOT_DIR / "aishua-master" / "src" / "main" / "resources" / "db"
COLUMN_LABEL_OVERRIDES = {
    "id": "id",
    "phone": "手机号",
    "nickname": "昵称",
    "password": "密码",
    "avatar": "头像地址",
    "status": "状态",
    "is_admin": "管理员标识",
    "create_time": "创建时间",
    "update_time": "更新时间",
    "deleted": "软删除",
    "name": "名称",
    "code": "编码",
    "description": "描述",
    "icon": "图标",
    "question_count": "题目总数",
    "sort": "排序",
    "is_enabled": "启用状态",
    "subject_id": "学科ID",
    "parent_id": "父目录ID",
    "tag": "备注",
    "title": "题干",
    "type": "题型",
    "directory_id": "目录ID",
    "difficulty": "难度",
    "answer": "答案",
    "analysis": "解析",
    "correct_rate": "正确率",
    "do_count": "做题次数",
    "session_code": "会话编码",
    "user_id": "用户ID",
    "practice_mode": "练习模式",
    "total_questions": "总题数",
    "answered_count": "已答题数",
    "correct_count": "正确题数",
    "wrong_count": "错误次数",
    "session_ref_id": "会话ID",
    "exercise_mode": "练习模式",
    "user_answer": "用户答案",
    "is_correct": "是否正确",
    "time_cost": "耗时",
    "exercise_time": "作答时间",
    "tags": "考点标签",
    "last_wrong_time": "最近错题时间",
    "master_status": "掌握状态",
    "analysis_code": "分析编码",
    "wrong_question_id": "错题ID",
    "question_id": "题目ID",
    "analysis_type": "分析类型",
    "result_json": "分析结果",
    "summary": "分析摘要",
    "model_name": "模型名称",
    "total_tokens": "总Token",
    "paper_name": "试卷名称",
    "total_score": "总分",
    "duration": "时长",
    "exam_name": "考试名称",
    "exam_mode": "考试模式",
    "correct_questions": "正确题数",
    "score": "得分",
}
ER_ENTITY_SPECS = [
    {
        "table": "user",
        "entity_name": "用户",
        "caption": "图4-4 用户实体属性图",
        "intro": "1. 用户实体：记录学生端与管理员端共用的账号基础信息。具体如图4-4所示。",
        "columns": ["id", "phone", "nickname", "password", "avatar", "status", "is_admin", "create_time", "update_time"],
        "width_cm": 12.8,
    },
    {
        "table": "subject",
        "entity_name": "学科",
        "caption": "图4-5 学科实体属性图",
        "intro": "2. 学科实体：记录题库所属学科及启用状态等基础信息。具体如图4-5所示。",
        "columns": ["id", "name", "code", "description", "question_count", "sort", "is_enabled", "create_time", "update_time"],
        "width_cm": 12.8,
    },
    {
        "table": "textbook_directory",
        "entity_name": "教材目录",
        "caption": "图4-6 教材目录实体属性图",
        "intro": "3. 教材目录实体：记录学科下的目录树层级与排序结构。具体如图4-6所示。",
        "columns": ["id", "name", "subject_id", "parent_id", "sort", "create_time", "update_time", "deleted"],
        "width_cm": 12.8,
    },
    {
        "table": "exam_tag",
        "entity_name": "考点标签",
        "caption": "图4-7 考点标签实体属性图",
        "intro": "4. 考点标签实体：记录学科知识点标签及其描述信息。具体如图4-7所示。",
        "columns": ["id", "name", "subject_id", "tag", "create_time", "update_time", "deleted"],
        "width_cm": 12.8,
    },
    {
        "table": "question",
        "entity_name": "题目",
        "caption": "图4-8 题目实体属性图",
        "intro": "5. 题目实体：记录题干、答案、解析、目录归属和难度等核心内容。具体如图4-8所示。",
        "columns": ["id", "title", "type", "subject_id", "directory_id", "difficulty", "answer", "analysis", "correct_rate", "do_count"],
        "width_cm": 12.8,
    },
    {
        "table": "practice_session",
        "entity_name": "练习会话",
        "caption": "图4-9 练习会话实体属性图",
        "intro": "6. 练习会话实体：记录一次刷题任务的整体过程状态。具体如图4-9所示。",
        "columns": ["id", "session_code", "user_id", "subject_id", "practice_mode", "question_count", "answered_count", "correct_count", "wrong_count", "status"],
        "width_cm": 12.8,
    },
    {
        "table": "exercise_record",
        "entity_name": "练习作答",
        "caption": "图4-10 练习作答实体属性图",
        "intro": "7. 练习作答实体：记录用户在练习过程中的单题作答明细。具体如图4-10所示。",
        "columns": ["id", "user_id", "question_id", "session_ref_id", "exercise_mode", "user_answer", "is_correct", "time_cost", "exercise_time"],
        "width_cm": 12.8,
    },
    {
        "table": "wrong_question",
        "entity_name": "错题",
        "caption": "图4-11 错题实体属性图",
        "intro": "8. 错题实体：记录用户错题沉淀、掌握状态和最近错误情况。具体如图4-11所示。",
        "columns": ["id", "user_id", "question_id", "subject_id", "directory_id", "tags", "wrong_count", "last_wrong_time", "master_status"],
        "width_cm": 12.8,
    },
    {
        "table": "wrong_question_ai_analysis",
        "entity_name": "错题AI分析",
        "caption": "图4-12 错题AI分析实体属性图",
        "intro": "9. 错题AI分析实体：记录错题诊断输入快照、分析结果与模型调用信息。具体如图4-12所示。",
        "columns": ["id", "analysis_code", "user_id", "wrong_question_id", "question_id", "analysis_type", "result_json", "summary", "model_name", "total_tokens"],
        "width_cm": 12.8,
    },
    {
        "table": "exam_paper",
        "entity_name": "试卷",
        "caption": "图4-13 试卷实体属性图",
        "intro": "10. 试卷实体：记录模拟考试试卷配置、总题数、总分和启用状态。具体如图4-13所示。",
        "columns": ["id", "paper_name", "subject_id", "total_questions", "total_score", "duration", "status", "create_time", "update_time"],
        "width_cm": 12.8,
    },
    {
        "table": "exam_record",
        "entity_name": "考试记录",
        "caption": "图4-14 考试记录实体属性图",
        "intro": "11. 考试记录实体：记录用户考试成绩、用时与完成状态。具体如图4-14所示。",
        "columns": ["id", "user_id", "subject_id", "exam_name", "exam_mode", "total_questions", "correct_questions", "score", "duration", "status"],
        "width_cm": 12.8,
    },
]
SCHEMA_TABLE_CACHE: dict[str, dict[str, dict[str, str]]] | None = None


def discover_template() -> Path:
    preferred = sorted(TEMPLATE_DIR.glob("*xxx*模板.docx"))
    if preferred:
        return preferred[0]
    candidates = sorted(TEMPLATE_DIR.glob("*模板.docx"))
    if not candidates:
        raise FileNotFoundError(f"未找到模板文件：{TEMPLATE_DIR}")
    return candidates[0]


def chinese_year_month(now: datetime) -> str:
    digits = {"0": "〇", "1": "一", "2": "二", "3": "三", "4": "四", "5": "五", "6": "六", "7": "七", "8": "八", "9": "九"}
    months = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七", 8: "八", 9: "九", 10: "十", 11: "十一", 12: "十二"}
    year = "".join(digits[ch] for ch in str(now.year))
    return f"{year}年{months[now.month]}月"


def set_run_font(run, size: float = 12, bold: bool = False, font_name: str = "宋体") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_paragraph(paragraph) -> None:
    paragraph.clear()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def set_paragraph_text(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    center: bool = False,
    indent: bool = False,
    size: float = 12,
    font_name: str = "宋体",
    line_spacing: float | None = 20,
) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if center else WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    paragraph.paragraph_format.line_spacing = Pt(line_spacing) if line_spacing else None
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=font_name)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    center: bool = False,
    indent: bool = False,
    style: str | None = None,
    size: float = 12,
    font_name: str = "宋体",
    line_spacing: float | None = 20,
):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    set_paragraph_text(
        paragraph,
        text,
        bold=bold,
        center=center,
        indent=indent,
        size=size,
        font_name=font_name,
        line_spacing=line_spacing,
    )
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    if text == "附录":
        add_paragraph(doc, text, center=True, size=16)
        return
    if re.match(r"^\d+\s", text):
        add_paragraph(doc, text, center=True, size=16)
        return
    if re.match(r"^\d+\.\d+\s", text):
        add_paragraph(doc, text, size=15)
        return
    if re.match(r"^\d+\.\d+\.\d+\s", text):
        add_paragraph(doc, text, size=14)
        return
    add_paragraph(doc, text)


def add_body(doc: Document, text: str) -> None:
    style = "Normal Indent" if "Normal Indent" in [item.name for item in doc.styles] else None
    add_paragraph(doc, text, indent=True, style=style)


def add_caption(doc: Document, text: str) -> None:
    style = "Caption" if "Caption" in [item.name for item in doc.styles] else None
    add_paragraph(doc, text, center=True, style=style)


def resolve_font_path(candidates: list[Path]) -> Path | None:
    for path in candidates:
        if path.exists():
            return path
    return None


def load_font(size: int, bold: bool = False):
    if ImageFont is None:
        return None
    font_path = resolve_font_path(BOLD_FONT_CANDIDATES if bold else REGULAR_FONT_CANDIDATES)
    if font_path is None:
        return ImageFont.load_default()
    return ImageFont.truetype(str(font_path), size)


def ensure_generated_figure_dir() -> None:
    GENERATED_FIGURE_DIR.mkdir(parents=True, exist_ok=True)


def discover_schema_sql() -> Path:
    preferred = sorted(SCHEMA_SQL_DIR.glob("*有数据*.sql"))
    if preferred:
        return preferred[0]
    candidates = sorted(SCHEMA_SQL_DIR.glob("*.sql"), key=lambda item: item.stat().st_size, reverse=True)
    if not candidates:
        raise FileNotFoundError(f"未找到数据库脚本文件：{SCHEMA_SQL_DIR}")
    return candidates[0]


def load_schema_table_cache() -> dict[str, dict[str, dict[str, str]]]:
    global SCHEMA_TABLE_CACHE
    if SCHEMA_TABLE_CACHE is not None:
        return SCHEMA_TABLE_CACHE
    schema_path = discover_schema_sql()
    tables: dict[str, dict[str, dict[str, str]]] = {}
    current_table: str | None = None
    for raw_line in schema_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip().rstrip(",")
        table_match = re.match(r"CREATE TABLE `([^`]+)`", line)
        if table_match:
            current_table = table_match.group(1)
            tables[current_table] = {}
            continue
        if current_table is None:
            continue
        if line.startswith(") ENGINE"):
            current_table = None
            continue
        if not line.startswith("`"):
            continue
        column_match = re.match(r"`([^`]+)`\s+([^\s,]+)(.*)", line)
        if column_match is None:
            continue
        comment_match = re.search(r"COMMENT\s+'([^']*)'", column_match.group(3))
        tables[current_table][column_match.group(1)] = {
            "type": column_match.group(2),
            "comment": comment_match.group(1) if comment_match else "",
        }
    SCHEMA_TABLE_CACHE = tables
    return tables


def shorten_comment_label(comment: str) -> str:
    if not comment:
        return ""
    label = re.sub(r"[（(].*?[）)]", "", comment)
    for separator in ("：", ":", "，", ",", "。", "；", ";"):
        if separator in label:
            label = label.split(separator, 1)[0]
    label = label.strip()
    if label == "主键ID":
        return "id"
    return label.replace("URL", "地址").replace("JSON", "结果").replace("统计", "")


def resolve_column_label(column_name: str, comment: str) -> str:
    if column_name in COLUMN_LABEL_OVERRIDES:
        return COLUMN_LABEL_OVERRIDES[column_name]
    shortened = shorten_comment_label(comment)
    if shortened:
        return shortened
    return column_name


def build_entity_attribute_items(spec: dict) -> list[dict[str, str | bool]]:
    tables = load_schema_table_cache()
    table_meta = tables.get(spec["table"], {})
    items: list[dict[str, str | bool]] = []
    for column_name in spec["columns"]:
        meta = table_meta.get(column_name, {})
        items.append(
            {
                "label": resolve_column_label(column_name, meta.get("comment", "")),
                "underline": column_name == "id",
            }
        )
    return items


def text_bbox(draw, text: str, font) -> tuple[int, int]:
    left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
    return right - left, bottom - top


def wrap_text(draw, text: str, font, max_width: int) -> list[str]:
    if not text:
        return [""]
    lines: list[str] = []
    for segment in text.splitlines() or [text]:
        if not segment:
            lines.append("")
            continue
        current = ""
        for char in segment:
            trial = current + char
            width, _ = text_bbox(draw, trial, font)
            if current and width > max_width:
                lines.append(current)
                current = char
            else:
                current = trial
        if current:
            lines.append(current)
    return lines or [text]


def draw_multiline_text_centered(draw, box: tuple[int, int, int, int], text: str, font, fill: str = "black", line_spacing: int = 6) -> None:
    x1, y1, x2, y2 = box
    max_width = max(10, x2 - x1 - 16)
    lines = wrap_text(draw, text, font, max_width)
    heights = [text_bbox(draw, line, font)[1] for line in lines]
    total_height = sum(heights) + line_spacing * max(0, len(lines) - 1)
    cursor_y = y1 + ((y2 - y1 - total_height) / 2)
    for line, height in zip(lines, heights):
        width, _ = text_bbox(draw, line, font)
        cursor_x = x1 + ((x2 - x1 - width) / 2)
        draw.text((cursor_x, cursor_y), line, font=font, fill=fill)
        cursor_y += height + line_spacing


def draw_box(
    draw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    *,
    width: int = 3,
    fill_color: str = "white",
    outline_color: str = "black",
    text_fill: str = "black",
) -> None:
    draw.rectangle(box, outline=outline_color, width=width, fill=fill_color)
    draw_multiline_text_centered(draw, box, text, font, fill=text_fill)


def draw_diamond(
    draw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    *,
    width: int = 3,
    fill_color: str = "white",
    outline_color: str = "black",
    text_fill: str = "black",
) -> None:
    x1, y1, x2, y2 = box
    points = [((x1 + x2) // 2, y1), (x2, (y1 + y2) // 2), ((x1 + x2) // 2, y2), (x1, (y1 + y2) // 2)]
    draw.polygon(points, outline=outline_color, fill=fill_color, width=width)
    draw_multiline_text_centered(draw, box, text, font, fill=text_fill)


def draw_ellipse_node(
    draw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    *,
    width: int = 3,
    fill_color: str = "white",
    outline_color: str = "black",
    text_fill: str = "black",
    underline: bool = False,
) -> None:
    draw.ellipse(box, outline=outline_color, width=width, fill=fill_color)
    if underline and "\n" not in text:
        text_width, text_height = text_bbox(draw, text, font)
        x = box[0] + ((box[2] - box[0] - text_width) / 2)
        y = box[1] + ((box[3] - box[1] - text_height) / 2)
        draw.text((x, y), text, font=font, fill=text_fill)
        underline_y = y + text_height + 4
        draw.line((x, underline_y, x + text_width, underline_y), fill=text_fill, width=2)
        return
    draw_multiline_text_centered(draw, box, text, font, fill=text_fill)


def draw_arrow(draw, start: tuple[int, int], end: tuple[int, int], *, width: int = 3, arrow_size: int = 10, fill: str = "black") -> None:
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    left = (
        x2 - arrow_size * math.cos(angle - math.pi / 6),
        y2 - arrow_size * math.sin(angle - math.pi / 6),
    )
    right = (
        x2 - arrow_size * math.cos(angle + math.pi / 6),
        y2 - arrow_size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([(x2, y2), left, right], fill=fill)


def draw_elbow_line(draw, points: list[tuple[int, int]], *, width: int = 3, fill: str = "black") -> None:
    for index in range(len(points) - 1):
        draw.line((*points[index], *points[index + 1]), fill=fill, width=width)


def draw_link(draw, start: tuple[int, int], end: tuple[int, int], *, width: int = 4, fill: str = "white") -> None:
    draw.line((start[0], start[1], end[0], end[1]), fill=fill, width=width)


def draw_text_badge(
    draw,
    position: tuple[int, int],
    text: str,
    font,
    *,
    fill_color: str = "white",
    text_fill: str = "black",
    boxed: bool = True,
) -> None:
    text_width, text_height = text_bbox(draw, text, font)
    x, y = position
    padding_x = 10
    padding_y = 5
    box = (x - padding_x, y - padding_y, x + text_width + padding_x, y + text_height + padding_y)
    if boxed:
        draw.rectangle(
            box,
            fill=fill_color,
        )
        draw.rectangle(box, outline="black", width=2)
    draw.text((x, y), text, font=font, fill=text_fill)


def draw_cardinality_labels(
    draw,
    start: tuple[int, int],
    end: tuple[int, int],
    font,
    *,
    start_card: str | None = None,
    end_card: str | None = None,
) -> None:
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = math.hypot(dx, dy) or 1
    nx = -dy / length
    ny = dx / length

    if start_card:
        ratio = 0.18
        x = int(start[0] + dx * ratio + nx * 18)
        y = int(start[1] + dy * ratio + ny * 18)
        draw_text_badge(draw, (x, y), start_card, font)
    if end_card:
        ratio = 0.82
        x = int(start[0] + dx * ratio + nx * 18)
        y = int(start[1] + dy * ratio + ny * 18)
        draw_text_badge(draw, (x, y), end_card, font)


def render_business_flow_figure(output_path: Path) -> Path:
    ensure_generated_figure_dir()
    width, height = 1800, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    box_font = load_font(30)
    title_font = load_font(38, bold=True)
    boxes = [
        ("用户登录\n与注册", (80, 300, 290, 430)),
        ("加入学科", (340, 300, 550, 430)),
        ("开始练习\n或考试", (600, 300, 860, 430)),
        ("提交作答", (910, 300, 1120, 430)),
        ("错题沉淀\n与统计更新", (1170, 300, 1430, 430)),
        ("AI分析\n与助教反馈", (1480, 300, 1690, 430)),
    ]
    for text, box in boxes:
        draw_box(draw, box, text, box_font)
    for idx in range(len(boxes) - 1):
        draw_arrow(draw, (boxes[idx][1][2], 365), (boxes[idx + 1][1][0], 365))
    draw_box(draw, (640, 80, 1160, 190), "持续强化学习闭环", title_font, width=4)
    draw_elbow_line(draw, [(1585, 300), (1585, 150), (185, 150), (185, 300)])
    draw_arrow(draw, (185, 150), (185, 300))
    image.save(output_path)
    return output_path


def render_data_flow_figure(output_path: Path) -> Path:
    ensure_generated_figure_dir()
    width, height = 1800, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    box_font = load_font(28)
    title_font = load_font(34, bold=True)
    left_box = (100, 360, 380, 520)
    right_top = (1320, 170, 1640, 310)
    right_middle = (1320, 400, 1640, 540)
    right_bottom = (1320, 630, 1640, 770)
    center_top = (670, 150, 1110, 290)
    center_middle = (670, 380, 1110, 520)
    center_bottom = (670, 610, 1110, 750)
    draw_box(draw, left_box, "学生端\n后台管理端", title_font, width=4)
    draw_box(draw, center_top, "请求接入层\n控制器与接口路由", box_font)
    draw_box(draw, center_middle, "核心业务处理层\n学科、练习、考试、AI分析", box_font)
    draw_box(draw, center_bottom, "数据持久化与状态缓存层", box_font)
    draw_box(draw, right_top, "题库与内容数据", box_font)
    draw_box(draw, right_middle, "作答记录与统计数据", box_font)
    draw_box(draw, right_bottom, "AI分析与外部资源", box_font)
    draw_arrow(draw, (380, 440), (670, 220))
    draw_arrow(draw, (890, 290), (890, 380))
    draw_arrow(draw, (890, 520), (890, 610))
    draw_arrow(draw, (1110, 220), (1320, 240))
    draw_arrow(draw, (1110, 450), (1320, 470))
    draw_arrow(draw, (1110, 680), (1320, 700))
    draw_arrow(draw, (1320, 240), (1110, 240))
    draw_arrow(draw, (1320, 470), (1110, 470))
    draw_arrow(draw, (1320, 700), (1110, 700))
    draw_arrow(draw, (670, 220), (380, 440))
    image.save(output_path)
    return output_path


def render_ipo_figure(output_path: Path) -> Path:
    ensure_generated_figure_dir()
    width, height = 1900, 1100
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, bold=True)
    box_font = load_font(24)
    small_font = load_font(22)
    draw_box(draw, (650, 40, 1250, 130), "关键功能IPO关系图", title_font, width=4)

    rows = [
        ("开始练习", "学科、目录、模式、题量", "校验身份并生成练习会话与题单", "返回题单与会话信息"),
        ("提交练习", "会话、作答、用时、图片", "判题并更新记录、错题与统计", "返回交卷结果与最新统计"),
        ("错题AI分析", "错题编号、上下文、分析请求", "读取错题快照并调用模型生成结果", "返回结构化分析与追问会话"),
        ("模拟考试", "试卷、作答、考试时间", "生成考试记录并汇总成绩", "返回成绩、明细与历史记录"),
    ]

    top_y = 190
    row_gap = 210
    for index, (name, input_text, process_text, output_text) in enumerate(rows):
        y = top_y + index * row_gap
        draw_box(draw, (70, y + 25, 280, y + 125), name, box_font, width=4)
        draw_box(draw, (380, y, 760, y + 150), f"输入\n{input_text}", small_font)
        draw_box(draw, (860, y, 1240, y + 150), f"处理\n{process_text}", small_font)
        draw_box(draw, (1340, y, 1720, y + 150), f"输出\n{output_text}", small_font)
        draw_arrow(draw, (280, y + 75), (380, y + 75))
        draw_arrow(draw, (760, y + 75), (860, y + 75))
        draw_arrow(draw, (1240, y + 75), (1340, y + 75))

    image.save(output_path)
    return output_path


def render_architecture_figure(output_path: Path) -> Path:
    ensure_generated_figure_dir()
    width, height = 1800, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    box_font = load_font(28)
    title_font = load_font(38, bold=True)
    draw_box(draw, (630, 70, 1170, 170), "AI刷题系统总体架构", title_font, width=4)
    draw_box(draw, (130, 260, 520, 420), "前端表示层\nVue3、Vue Router\nVuex、Vant、ECharts", box_font)
    draw_box(draw, (700, 260, 1100, 420), "后端业务层\nSpringBoot 3.1.8\nMyBatis-Plus、JWT", box_font)
    draw_box(draw, (1280, 210, 1660, 330), "MySQL\n核心业务数据", box_font)
    draw_box(draw, (1280, 390, 1660, 510), "Redis\n练习草稿与短时状态", box_font)
    draw_box(draw, (1280, 570, 1660, 690), "MinIO\n题图与作答图片", box_font)
    draw_box(draw, (1280, 750, 1660, 870), "AI服务\n千问兼容接口\nDify 工作流", box_font)
    draw_arrow(draw, (520, 340), (700, 340))
    for y in (270, 450, 630, 810):
        draw_arrow(draw, (1100, 340), (1280, y))
        draw_arrow(draw, (1280, y), (1100, 340))
    image.save(output_path)
    return output_path


def render_module_figure(output_path: Path) -> Path:
    ensure_generated_figure_dir()
    width, height = 2100, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    center_font = load_font(38, bold=True)
    box_font = load_font(28)
    parent = (800, 80, 1300, 200)
    children = [
        ("用户与权限\n模块", (90, 520, 330, 690)),
        ("学科与目录\n模块", (380, 520, 620, 690)),
        ("练习闭环\n模块", (670, 520, 910, 690)),
        ("AI智能分析\n模块", (960, 520, 1200, 690)),
        ("模拟考试\n模块", (1250, 520, 1490, 690)),
        ("学情统计\n模块", (1540, 520, 1780, 690)),
        ("文件资源\n模块", (1830, 520, 2070, 690)),
    ]
    draw_box(draw, parent, "AI刷题系统", center_font, width=4)
    parent_bottom_x = (parent[0] + parent[2]) // 2
    parent_bottom_y = parent[3]
    bus_y = 360
    draw_elbow_line(draw, [(parent_bottom_x, parent_bottom_y), (parent_bottom_x, bus_y)])
    draw.line((90, bus_y, 2070, bus_y), fill="black", width=3)
    for text, box in children:
        draw_box(draw, box, text, box_font)
        center_x = (box[0] + box[2]) // 2
        draw.line((center_x, bus_y, center_x, box[1]), fill="black", width=3)
    image.save(output_path)
    return output_path


def render_er_figure(output_path: Path) -> Path:
    ensure_generated_figure_dir()
    width, height = 1800, 1260
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    entity_font = load_font(34)
    relation_font = load_font(30)
    cardinality_font = load_font(36, bold=True)
    legend_font = load_font(24)
    entity_boxes = {
        "学科": (90, 160, 310, 235),
        "教材目录": (90, 430, 310, 505),
        "考点标签": (90, 700, 310, 775),
        "题目": (610, 430, 830, 505),
        "试卷": (610, 700, 830, 775),
        "用户": (980, 80, 1200, 155),
        "练习会话": (980, 330, 1200, 405),
        "练习作答": (1360, 330, 1580, 405),
        "错题记录": (1360, 590, 1580, 665),
        "AI分析": (1360, 850, 1580, 925),
        "考试记录": (980, 940, 1200, 1015),
    }
    relation_boxes = {
        "加入": {"box": (520, 95, 660, 215), "label": "加入"},
        "包含": {"box": (140, 295, 260, 405), "label": "包含"},
        "关联": {"box": (140, 565, 260, 675), "label": "关联"},
        "收录": {"box": (360, 415, 500, 525), "label": "收录"},
        "标注": {"box": (360, 665, 500, 775), "label": "标注"},
        "组卷": {"box": (660, 565, 780, 675), "label": "组卷"},
        "发起": {"box": (1030, 205, 1150, 315), "label": "发起"},
        "生成练习": {"box": (1210, 315, 1330, 425), "label": "生成"},
        "对应": {"box": (1010, 480, 1130, 590), "label": "对应"},
        "沉淀": {"box": (1410, 455, 1530, 565), "label": "沉淀"},
        "诊断": {"box": (1410, 715, 1530, 825), "label": "诊断"},
        "参加": {"box": (880, 560, 1000, 670), "label": "参加"},
        "生成考试": {"box": (760, 830, 880, 940), "label": "生成"},
    }

    def center(box: tuple[int, int, int, int]) -> tuple[int, int]:
        return ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)

    connectors = [
        ("学科", "加入"),
        ("加入", "用户"),
        ("学科", "包含"),
        ("包含", "教材目录"),
        ("教材目录", "关联"),
        ("关联", "考点标签"),
        ("教材目录", "收录"),
        ("收录", "题目"),
        ("考点标签", "标注"),
        ("标注", "题目"),
        ("题目", "组卷"),
        ("组卷", "试卷"),
        ("用户", "发起"),
        ("发起", "练习会话"),
        ("练习会话", "生成练习"),
        ("生成练习", "练习作答"),
        ("题目", "对应"),
        ("对应", "练习作答"),
        ("练习作答", "沉淀"),
        ("沉淀", "错题记录"),
        ("错题记录", "诊断"),
        ("诊断", "AI分析"),
        ("用户", "参加"),
        ("参加", "考试记录"),
        ("试卷", "生成考试"),
        ("生成考试", "考试记录"),
    ]
    all_boxes = {**entity_boxes, **{name: spec["box"] for name, spec in relation_boxes.items()}}
    for start_name, end_name in connectors:
        draw_link(draw, center(all_boxes[start_name]), center(all_boxes[end_name]), width=5, fill="black")
    for text, box in entity_boxes.items():
        draw_box(draw, box, text, entity_font, width=4, fill_color="#F8F8F8", outline_color="black", text_fill="black")
    for _, spec in relation_boxes.items():
        draw_diamond(draw, spec["box"], spec["label"], relation_font, width=4, fill_color="white", outline_color="black", text_fill="black")
    cardinality_labels = [
        ("N", (318, 174)), ("M", (950, 90)),
        ("1", (180, 246)), ("N", (182, 408)),
        ("M", (180, 516)), ("N", (180, 672)),
        ("1", (318, 448)), ("N", (580, 448)),
        ("M", (318, 720)), ("N", (590, 505)),
        ("M", (735, 515)), ("N", (670, 790)),
        ("1", (1068, 158)), ("N", (1072, 310)),
        ("1", (1205, 330)), ("N", (1325, 322)),
        ("1", (835, 468)), ("N", (1305, 410)),
        ("N", (1475, 410)), ("1", (1480, 565)),
        ("1", (1482, 668)), ("N", (1484, 825)),
        ("1", (945, 160)), ("N", (1030, 910)),
        ("1", (835, 740)), ("N", (945, 975)),
    ]
    for text, position in cardinality_labels:
        draw_text_badge(draw, position, text, cardinality_font, boxed=False)
    draw_text_badge(draw, (1320, 1125), "1、N、M 表示实体间联系基数", legend_font, boxed=False)
    image.save(output_path)
    return output_path


def render_entity_attribute_figure(output_path: Path, entity_name: str, attributes: list[dict[str, str | bool]]) -> Path:
    ensure_generated_figure_dir()
    width, height = 1300, 1300
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    entity_font = load_font(34, bold=True)
    attr_font = load_font(28)
    entity_box = (500, 565, 800, 695)
    entity_center = ((entity_box[0] + entity_box[2]) / 2, (entity_box[1] + entity_box[3]) / 2)
    radius = 430
    ellipse_width = 230
    ellipse_height = 120
    trimmed_attributes = attributes[:10]
    for index, item in enumerate(trimmed_attributes):
        angle = math.radians(-90 + (360 / len(trimmed_attributes)) * index)
        attr_center_x = entity_center[0] + radius * math.cos(angle)
        attr_center_y = entity_center[1] + radius * math.sin(angle)
        box = (
            int(attr_center_x - ellipse_width / 2),
            int(attr_center_y - ellipse_height / 2),
            int(attr_center_x + ellipse_width / 2),
            int(attr_center_y + ellipse_height / 2),
        )
        draw_link(draw, (int(entity_center[0]), int(entity_center[1])), (int(attr_center_x), int(attr_center_y)), width=5, fill="black")
        draw_ellipse_node(
            draw,
            box,
            str(item["label"]),
            attr_font,
            width=4,
            fill_color="white",
            outline_color="black",
            text_fill="black",
            underline=bool(item["underline"]),
        )
    draw_box(draw, entity_box, entity_name, entity_font, width=4, fill_color="#F3F3F3", outline_color="black", text_fill="black")
    image.save(output_path)
    return output_path


def make_entity_attribute_renderer(spec: dict):
    def renderer(output_path: Path, current_spec: dict = spec) -> Path:
        return render_entity_attribute_figure(output_path, current_spec["entity_name"], build_entity_attribute_items(current_spec))

    return renderer


AUTO_FIGURE_BUILDERS = {
    "图3-1 学生端核心业务流程图": ("figure_3_1_business_flow.png", render_business_flow_figure),
    "图3-2 系统主要业务数据流图": ("figure_3_2_data_flow.png", render_data_flow_figure),
    "图3-3 关键功能IPO图": ("figure_3_3_ipo.png", render_ipo_figure),
    "图4-1 系统总体架构图": ("figure_4_1_architecture.png", render_architecture_figure),
    "图4-2 系统功能模块图": ("figure_4_2_modules.png", render_module_figure),
    "图4-3 总体E-R图": ("figure_4_3_er.png", render_er_figure),
}
for spec in ER_ENTITY_SPECS:
    AUTO_FIGURE_BUILDERS[spec["caption"]] = (f"figure_{spec['table']}_er.png", make_entity_attribute_renderer(spec))


CAPTURED_FIGURES = {
    1: "figure_5_1_ui.png",
    2: "figure_5_2_code.png",
    3: "figure_5_3_ui.png",
    4: "figure_5_4_code.png",
    5: "figure_5_5_ui.png",
    6: "figure_5_6_code.png",
    7: "figure_5_7_ui.png",
    8: "figure_5_8_code.png",
    10: "figure_5_10_code.png",
    12: "figure_5_12_code.png",
    13: "figure_5_13_ui.png",
    14: "figure_5_14_code.png",
}


def build_figure(caption: str) -> Path | None:
    figure_match = re.match(r"^\u56fe5-(\d+)\s", caption)
    if figure_match:
        filename = CAPTURED_FIGURES.get(int(figure_match.group(1)))
        if filename:
            captured_path = CAPTURED_FIGURE_DIR / filename
            if captured_path.exists():
                return captured_path

    if Image is None:
        return None
    spec = AUTO_FIGURE_BUILDERS.get(caption)
    if spec is None:
        return None
    filename, renderer = spec
    output_path = GENERATED_FIGURE_DIR / filename
    return renderer(output_path)


def add_capture_guide(doc: Document, text: str) -> None:
    add_paragraph(doc, text, size=10.5, line_spacing=16)


def add_figure_placeholder(doc: Document, caption: str, note: str, guide: str | None = None, width_cm: float = 15.8) -> None:
    add_caption(doc, caption)
    figure_path = build_figure(caption)
    if figure_path is not None and figure_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(str(figure_path), width=Cm(width_cm))
    else:
        add_paragraph(doc, f"【插图位置：{note}】", center=True)
        if guide:
            add_capture_guide(doc, guide)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.first_line_indent = None
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, font_name="Consolas")
        if index != len(lines) - 1:
            run.add_break()


def set_cell_width(cell, width_cm: float) -> None:
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if hasattr(table, "autofit"):
        table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_table(
    doc: Document,
    rows: list[list[str]],
    *,
    column_widths: list[float] | None = None,
    font_size: float = 12,
    header_font_size: float | None = None,
    line_spacing: float | None = 16,
    column_alignments: list[str] | None = None,
) -> None:
    if not rows or not rows[0]:
        return
    column_count = len(rows[0])
    if column_widths is not None and len(column_widths) != column_count:
        raise ValueError("列宽数量与表格列数不一致")
    if column_alignments is not None and len(column_alignments) != column_count:
        raise ValueError("列对齐配置数量与表格列数不一致")

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_fixed_layout(table)
    header_font_size = header_font_size or font_size

    header = table.rows[0].cells
    for index, cell_value in enumerate(rows[0]):
        if column_widths is not None:
            set_cell_width(header[index], column_widths[index])
        header[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_paragraph_text(
            header[index].paragraphs[0],
            cell_value,
            bold=True,
            center=True,
            size=header_font_size,
            line_spacing=line_spacing,
        )
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for index, cell_value in enumerate(row_values):
            if column_widths is not None:
                set_cell_width(cells[index], column_widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            align = column_alignments[index] if column_alignments is not None else "left"
            set_paragraph_text(
                cells[index].paragraphs[0],
                cell_value,
                center=align == "center",
                size=font_size,
                line_spacing=line_spacing,
            )


def find_paragraph_index(doc: Document, text: str, start: int = 0) -> int:
    for idx in range(start, len(doc.paragraphs)):
        if doc.paragraphs[idx].text.strip() == text:
            return idx
    raise ValueError(f"未找到段落：{text}")


def remove_extra_abstract_paragraphs(doc: Document, abstract_idx: int, keyword_idx: int) -> None:
    body = doc._element.body
    keep_element = doc.paragraphs[abstract_idx]._element
    keep_index = list(body).index(keep_element)
    end_element = doc.paragraphs[keyword_idx]._element
    end_index = list(body).index(end_element)
    for child in list(body)[keep_index + 1:end_index]:
        if child.tag.endswith("}p"):
            body.remove(child)


def patch_front_matter(doc: Document) -> None:
    if doc.tables:
        doc.tables[0].cell(0, 1).text = TITLE
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            if row.cells[1].text.strip() == "":
                row.cells[1].text = "待填写"
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped.startswith("二〇") and "年" in stripped and "月" in stripped:
            set_paragraph_text(paragraph, chinese_year_month(datetime.now()), center=True)
            break
    abstract_title_idx = find_paragraph_index(doc, "摘要")
    keyword_idx = next(idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip().startswith("关键词"))
    abstract_idx = abstract_title_idx + 1
    while abstract_idx < keyword_idx and not doc.paragraphs[abstract_idx].text.strip():
        abstract_idx += 1
    set_paragraph_text(doc.paragraphs[abstract_idx], ABSTRACT, indent=True)
    set_paragraph_text(doc.paragraphs[keyword_idx], KEYWORDS)
    remove_extra_abstract_paragraphs(doc, abstract_idx, keyword_idx)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in TOC_REPLACEMENTS:
            set_paragraph_text(paragraph, TOC_REPLACEMENTS[paragraph.text.strip()])


def rebuild_body(doc: Document) -> None:
    start_idx = find_paragraph_index(doc, "1 引言")
    body = doc._element.body
    start_element = doc.paragraphs[start_idx]._element
    start_pos = list(body).index(start_element)
    for child in list(body)[start_pos:]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "1 引言")
    add_heading(doc, "1.1 项目背景及概述")
    add_heading(doc, "1.1.1 项目背景")
    add_body(doc, "随着在线教育平台、移动学习工具和大模型技术的快速发展，学习者对刷题系统的期待已经由单纯的题目练习转向更完整的学习支持。传统刷题平台通常只强调题目展示、答案判定与分数统计，但在知识点组织、错题沉淀、学习诊断和个性化反馈等方面仍存在明显不足。")
    add_body(doc, "从教育信息化的发展趋势来看，题库系统已经不再只是提供练习资源的工具，而应逐渐成为连接学习过程、学习数据和智能分析能力的载体。尤其在课程训练、高中复习、资格考试与阶段性测验等场景中，学习者更加需要清晰的目录结构、持续的错题跟踪和针对性的强化建议。")
    add_body(doc, "本项目正是在这一背景下提出，目标是设计并实现一套基于SpringBoot与Vue3的AI刷题系统，使系统既具备传统题库平台的题目管理与考试能力，又具备错题智能分析、学习画像和薄弱点推荐等增强能力。")

    add_heading(doc, "1.1.2 问题概述")
    add_body(doc, "本文主要讨论并解决的问题，是如何构建一套围绕学习闭环运行的AI刷题系统。现有不少题库平台虽然能够提供基本的练习功能，但题目与目录、考点之间的关系不够清晰，导致后续按章节学习、按知识点强化和按学科统计时缺乏稳定支撑。")
    add_body(doc, "同时，很多系统在交卷后只给出对错和分数，无法把练习结果真正沉淀为可持续利用的错题资源和学情数据，学习者虽然知道自己做错了，却无法准确定位问题原因，也难以获得下一步的强化建议。")
    add_body(doc, "在AI能力方面，一些系统仅把大模型作为独立问答窗口使用，没有与错题、作答上下文和统计数据结合，因此无法形成真正有针对性的学习反馈。基于这些问题，本文围绕题目组织、练习闭环、错题复盘、考试流程和智能分析能力的协同展开研究。")

    add_heading(doc, "1.2 研究的意义和重要性")
    add_heading(doc, "1.2.1 研究的意义")
    add_body(doc, "从理论层面看，本文把题库系统、学习过程数据和AI能力纳入同一研究对象，不再把刷题系统简单理解为题目展示工具，而是将其视作一套围绕学习闭环持续运行的数据系统。这种视角有助于补充传统课程练习平台在错题诊断、掌握度评估和持续反馈方面的研究思路。")
    add_body(doc, "从实践层面看，本文所设计的系统与1.1.2小节提出的问题形成直接呼应。系统通过学科、目录和考点结构化组织解决题目资源分散的问题，通过练习会话、错题记录和统计数据解决练习结果难以沉淀的问题，通过AI分析和资源推荐提升反馈深度，因此具有较强的实际应用价值。")

    add_heading(doc, "1.2.2 研究的重要性")
    add_body(doc, "本课题的重要性主要体现在两个方面。第一，它验证了大模型能力与传统题库业务并非替代关系，而是可以在完整数据支撑下形成协同，从而提升学习诊断与强化效果。")
    add_body(doc, "第二，它为后续扩展个性化补题、学习路径推荐、智能组卷和课程画像等能力奠定了数据和流程基础。本文所解决的，是智能学习平台中更底层、更关键的支撑问题，因此具有持续演进价值。")

    add_heading(doc, "1.3 研究内容和主要工作")
    add_heading(doc, "1.3.1 研究内容")
    add_body(doc, "本文围绕AI刷题系统的需求分析、系统设计、功能实现和部署验证展开研究。系统采用前后端分离架构，后端使用SpringBoot 3.1.8、MyBatis-Plus、JWT、Redis和MinIO构建业务服务与基础设施能力，前端基于Vue3、Vue Router、Vuex、Vant和ECharts完成学生端与后台管理端页面实现，智能能力则通过千问兼容接口和Dify工作流接入。")
    add_body(doc, "在系统组织结构上，项目以学科为主线，以目录与考点为资源组织层，以题库、练习、错题、考试和学情统计为核心业务层，以AI分析能力为增强层，形成“资源组织、过程记录、结果沉淀、智能反馈”四层协同结构。")
    add_body(doc, "本文研究并实现了学科管理、目录与考点组织、练习会话与草稿缓存、错题记录与智能分析、模拟考试流程、学习统计与薄弱点推荐，以及本地混合部署方案整理等内容。")

    add_heading(doc, "1.3.2 主要工作")
    add_body(doc, "围绕上述研究内容，本文作者本人主要完成了以下工作：第一，结合当前项目代码与数据库脚本，梳理了学生端与后台管理端的业务流程，明确系统的主要参与者、核心用例与数据流向。")
    add_body(doc, "第二，完成了系统总体结构与数据库关系的整理工作，对用户、学科、目录、考点、题目、练习、考试和AI分析等主要实体及其关系进行了归纳。")
    add_body(doc, "第三，完成了对后端主要模块实现方式的梳理，包括JWT鉴权、练习草稿缓存、批量交卷、错题AI分析会话持久化、试卷配置、考试记录和文件上传等内容。")
    add_body(doc, "第四，完成了对前端功能结构和页面职责的归纳，并根据模板要求撰写需求、设计、实现、部署和总结等章节内容。")
    add_body(doc, "第五，如果将本项目视为课程综合实践项目中的整体成果，则本文作者本人承担的是需求分析、设计说明、功能梳理、部署验证方案整理以及文档撰写等主要工作。")

    add_heading(doc, "1.4 论文结构")
    add_body(doc, "本文共分为七个章节。第一章介绍项目背景、问题概述、研究意义、研究内容以及作者本人所承担的主要工作。第二章说明系统开发方法与关键技术。第三章从业务需求、功能需求、数据需求和非功能需求四个方面展开系统需求分析。")
    add_body(doc, "第四章描述系统总体架构、功能模块和数据库设计。第五章结合项目当前代码说明主要功能模块的实现思路。第六章给出系统在Windows宿主机与Docker Desktop混合环境下的部署规划、操作步骤和验证方法。第七章对全文进行总结，并提出未来完善方向。")

    add_heading(doc, "2 系统开发方法及相关技术概述")
    add_body(doc, "本章先从整体上概述系统开发方法和技术选型，再分别说明各项关键技术在系统中的具体作用。由于本文所研究的是一个集题库组织、练习过程记录、考试管理和AI分析于一体的综合系统，因此技术选择既要考虑开发效率，也要兼顾结构清晰和后续扩展。")
    add_heading(doc, "2.1 系统开发方法")
    add_body(doc, "本系统采用原型迭代与前后端分离相结合的开发方法。首先围绕“学科、练习、错题、考试”主链路构建最小可运行版本，确保用户能够完成从登录到学习反馈的基本闭环；在核心流程稳定后，再逐步补充AI错题分析、主观题辅助评分、对象存储和外部工作流等增强能力。")
    add_body(doc, "在实现层面，后端按业务域进行模块划分，控制层负责请求接入和参数校验，服务层负责流程编排，数据访问层负责业务数据读写，公共支撑层承载权限、上下文和基础设施封装。前端则以页面和模块为单位组织代码，配合Vue Router与Vuex完成路由和状态管理。")

    add_heading(doc, "2.2 相关技术概述")
    add_heading(doc, "2.2.1 SpringBoot框架概述")
    add_body(doc, "SpringBoot是系统后端的核心开发框架，负责应用启动、配置装配、Web接口暴露以及依赖整合。项目当前基于SpringBoot 3.1.8和Java 17实现，配合Jakarta Validation、统一结果封装和拦截器机制组织接口层结构。")
    add_body(doc, "在本系统中，SpringBoot主要承担用户登录注册、练习流程控制、考试管理、错题AI分析、文件上传和后台管理等核心业务。")

    add_heading(doc, "2.2.2 Vue3框架概述")
    add_body(doc, "Vue3用于构建系统前端页面，负责学生端与后台管理端的交互实现。项目通过Vue Router组织页面跳转，使用Vuex维护登录状态与用户信息，并配合Vant和ECharts完成组件交互与学习数据展示。")
    add_body(doc, "相较于传统静态页面，Vue3更适合本系统这种模块较多、交互频繁、状态切换明显的业务场景。")

    add_heading(doc, "2.2.3 MySQL数据库概述")
    add_body(doc, "MySQL是系统的主要业务数据库，承担用户、学科、目录、考点、题目、练习、考试和AI分析等核心数据的持久化任务。当前SQL脚本中已定义多张核心业务表，能够支撑题库组织、作答记录和统计分析等需求。")

    add_heading(doc, "2.2.4 Redis概述")
    add_body(doc, "Redis在本系统中主要用于练习草稿缓存和短时状态存储。由于练习作答过程中的数据变化较为频繁，若每次输入都直接写入数据库，会对交互体验产生不利影响，因此项目通过Redis暂存草稿并按策略同步。")

    add_heading(doc, "2.2.5 MinIO对象存储概述")
    add_body(doc, "MinIO用于管理题图、主观题作答图片和其他对象文件资源。与直接把大体积文件存放在数据库中相比，对象存储更有利于提升文件管理效率，并保持业务数据结构清晰。")

    add_heading(doc, "2.2.6 Dify与大模型接口概述")
    add_body(doc, "系统智能能力采用两类接入方式：一类是直接对接千问兼容接口，用于错题分析、练习题智能助教和主观题辅助评分等实时性较强的场景；另一类是通过Dify工作流完成视频检索等适合流程编排的任务。")
    add_body(doc, "这种方式兼顾了低延迟问答与可扩展工作流两方面需求，使AI能力能够更自然地嵌入学习流程。")

    add_heading(doc, "3 系统需求分析")
    add_body(doc, "需求分析的目的是明确系统需要解决的主要问题、系统应具备的功能边界以及相关运行要求，为后续设计和实现提供依据。结合当前项目代码与模板要求，本章从业务需求、功能需求、数据需求和非功能需求四个方面展开分析。")

    add_heading(doc, "3.1 业务需求")
    add_body(doc, "从技术可行性看，本项目采用的SpringBoot、Vue3、MySQL、Redis、MinIO和Dify等技术方案都较为成熟，且当前项目已经具备本地运行基础，因此在实现、联调和课程展示层面均具有较强可行性。")
    add_body(doc, "从学生端角度看，当前主要业务问题在于学习过程缺乏连续性。很多题库平台只能完成“选题、作答、看结果”的基础操作，错题难以有效归集，统计结果难以形成学情画像，薄弱点也无法自动转化为后续强化内容。")
    add_body(doc, "从管理端角度看，当前主要业务问题在于学习资源维护分散。如果系统缺少统一的学科、目录、考点、题库和试卷管理能力，那么学生端即使有练习页面，也很难获得稳定、规范、可分析的学习内容。")
    add_body(doc, "综合以上分析，本系统需要达到的主要业务目标包括：建立完整的学科练习闭环、让练习结果自动沉淀为错题与统计数据、让后台具备稳定的内容治理能力、让AI能力建立在真实学习上下文之上。")
    add_figure_placeholder(doc, "图3-1 学生端核心业务流程图", "此处插入学生端业务流程图，可展示“登录—加入学科—开始练习—提交作答—错题沉淀—AI分析—持续强化”的流程。")
    add_body(doc, "学生端的主要业务流程可以概括为：用户登录后选择并加入学科，浏览目录和考点，发起练习或考试，系统在交卷后自动生成作答记录、错题记录和统计结果，随后用户可以继续查看学情分析并通过AI能力完成错题复盘与强化。")

    add_heading(doc, "3.2 功能需求")
    add_heading(doc, "3.2.1 角色分析")
    add_body(doc, "从系统角度看，AI刷题系统的参与者并非只有学生用户，还包括管理员、AI分析服务以及底层存储与缓存服务。各参与者职责边界清晰，才能保证系统接口设计和模块划分稳定合理。")
    add_caption(doc, "表3-1 角色分析表")
    add_table(
        doc,
        [
            ["角色", "职责描述"],
            ["学生用户", "注册登录、加入学科、浏览目录与考点、发起练习与考试、查看错题、查看学情分析并触发AI分析。"],
            ["管理员", "维护学科、教材目录、考点标签、题库、试卷和考试记录，保障平台内容供给和后台运营。"],
            ["AI分析服务", "基于错题、作答上下文与统计数据输出错因分析、学习建议、视频检索结果和辅助评分结果。"],
            ["数据与文件服务", "通过MySQL、Redis和MinIO分别承担业务持久化、草稿缓存和对象文件管理，为系统运行提供底层支撑。"],
        ],
    )
    add_body(doc, "由表3-1可以看出，学生用户是核心服务对象，管理员负责内容治理，AI分析服务负责增强反馈能力，基础数据与文件服务负责保证整体运行稳定。")

    add_heading(doc, "3.2.2 业务分析")
    add_body(doc, "从系统使用者角度分析，学生端功能主要围绕学科、练习、错题、考试和学情展开，管理员端则围绕学科、目录、考点、题库和试卷维护展开。为了更好地描述这些功能之间的数据流转关系，需要通过数据流图对系统主要业务进行表达。")
    add_figure_placeholder(doc, "图3-2 系统主要业务数据流图", "此处插入系统数据流图，可体现学生端、后台管理端、题库数据、统计数据和AI分析之间的流转关系。")
    add_body(doc, "在学生端数据流中，用户登录后首先获取身份信息与学科权限，再进入目录与考点浏览、练习题目获取、交卷判题、错题沉淀和学习统计等流程；在后台管理端数据流中，管理员通过学科、目录、题目和试卷维护接口持续为前台学习闭环提供可用内容。")
    add_caption(doc, "表3-2 核心数据字典")
    add_table(
        doc,
        [
            ["数据项", "来源", "说明"],
            ["user_id", "JWT解析得到的当前登录用户", "用于标识练习、考试、错题和AI会话归属。"],
            ["subject_id", "学科选择或已加入学科记录", "用于标识目录、题目、统计和考试的学科范围。"],
            ["directory_id", "教材目录节点", "用于定位题目所属章节，并支撑按目录练习与统计。"],
            ["tag_id / tag_ids", "考点标签关系数据", "用于描述题目知识点及薄弱点分析依据。"],
            ["session_id", "练习会话或AI会话编号", "用于串联一次完整的练习或问答过程。"],
            ["wrong_question_id", "错题记录主键", "用于关联错题本、AI分析和追问会话。"],
            ["paper_id / exam_record_id", "试卷与考试记录主键", "用于支撑模拟考试流程。"],
            ["analysis_code", "错题AI分析唯一编码", "用于标识一次结构化分析结果快照。"],
        ],
    )

    add_heading(doc, "3.3 数据需求")
    add_body(doc, "系统的数据需求应与3.2.2小节中的业务分析相对应。结合当前项目实现，系统需要管理基础资源数据、过程性业务数据、统计画像数据以及AI分析数据四大类信息。")
    add_body(doc, "其中，基础资源数据包括学科、目录、考点与题目；过程性业务数据包括练习会话、作答记录、错题记录、试卷与考试记录；统计画像数据包括用户统计、学科统计、每日统计和知识点掌握度；AI分析数据包括错题分析快照、追问会话和消息记录。")
    add_body(doc, "为了说明系统关键功能在输入、处理和输出方面的关系，本文进一步以开始练习、提交练习、错题AI分析和模拟考试为代表给出IPO分析。")
    add_figure_placeholder(doc, "图3-3 关键功能IPO图", "此处插入关键功能IPO图，可围绕开始练习、提交练习、错题AI分析和模拟考试展示输入、处理与输出之间的关系。")
    add_caption(doc, "表3-3 核心功能IPO分析表")
    add_table(
        doc,
        [
            ["功能", "输入", "处理", "输出"],
            ["开始练习", "subject_id、directory_id、练习模式、题量等参数", "校验用户状态并生成练习会话和题单", "练习会话信息与题目列表。"],
            ["提交练习", "session_id、作答内容、用时和图片等", "判题并更新作答记录、错题记录和统计数据", "交卷结果、对错统计和最新学习数据。"],
            ["错题AI分析", "wrong_question_id、分析请求和用户上下文", "读取错题、题目、最近作答记录和掌握度并调用模型生成分析结果", "结构化分析结果与可继续追问的会话。"],
            ["模拟考试", "paper_id、用户作答、考试时间", "生成考试记录并完成成绩汇总", "考试总成绩、题目明细和考试历史。"],
        ],
    )

    add_heading(doc, "3.4 非功能需求")
    add_heading(doc, "3.4.1 环境需求")
    add_body(doc, "系统环境需求应区分硬件环境和软件环境。由于当前项目主要面向本地开发与课程展示，因此硬件方面重点强调在Windows宿主机与Docker Desktop混合环境下的稳定运行能力，软件方面则要求各基础运行时与项目依赖版本保持一致。")
    add_caption(doc, "表3-4 硬件环境需求")
    add_table(
        doc,
        [
            ["类别", "建议配置", "说明"],
            ["开发与部署宿主机", "64位CPU、8GB及以上内存、50GB及以上可用磁盘", "用于运行前端、后端、本地数据库客户端和Docker Desktop。"],
            ["网络环境", "可访问localhost相关服务", "满足前端访问后端以及后端访问中间件和AI服务的需求。"],
            ["容器运行能力", "支持Docker Desktop", "用于运行Redis、MinIO和Dify等基础服务。"],
        ],
    )
    add_caption(doc, "表3-5 软件环境需求")
    add_table(
        doc,
        [
            ["类别", "名称", "说明"],
            ["操作系统", "Windows 10/11 宿主机", "项目当前主要运行平台。"],
            ["JDK版本", "JDK 17", "满足SpringBoot 3.1.8运行要求。"],
            ["后端框架", "SpringBoot 3.1.8 + MyBatis-Plus", "负责业务服务与数据访问。"],
            ["前端框架", "Vue3 + Vue Router + Vuex + Vant + ECharts", "负责学生端与后台管理端页面。"],
            ["数据库", "MySQL 8.0", "保存核心业务数据。"],
            ["缓存", "Redis", "保存练习草稿与短时状态。"],
            ["对象存储", "MinIO", "管理题图和作答图片等文件。"],
            ["AI服务", "千问兼容接口 + Dify", "提供错题分析、问答和资源检索能力。"],
        ],
    )

    add_heading(doc, "3.4.2 安全性需求")
    add_body(doc, "系统应保证账号、作答数据和后台管理功能的安全性。当前后端通过JWT进行登录态识别，并在AuthInterceptor中对/api路径进行统一拦截，仅放行登录和注册接口。")
    add_body(doc, "对于后台管理接口，系统通过is_admin字段和AdminAccess校验权限边界，防止普通用户访问管理功能。对于文件访问和AI分析会话，也应确保其归属关系明确，避免用户越权读取他人数据。")

    add_heading(doc, "3.4.3 用户体验需求")
    add_body(doc, "从学生端看，系统应保证练习链路简洁顺畅，用户能够快速完成加入学科、浏览目录、开始练习、保存草稿和查看结果等操作。对于错题和AI分析场景，系统应尽量减少重复输入，让用户基于已有记录直接发起复盘。")
    add_body(doc, "从管理端看，列表筛选、表单录入和试卷配置应保持清晰结构，避免因字段过多导致操作成本过高。统计页面应突出对学习过程真正有帮助的信息。")

    add_heading(doc, "3.4.4 性能需求")
    add_body(doc, "系统性能需求主要体现在以下几个方面：普通查询接口应保持较快响应速度，练习草稿保存不能明显打断答题节奏，批量交卷时应保证多表更新的一致性，分页列表在数据量增多后仍需稳定返回结果。")
    add_body(doc, "对于AI流式问答场景，系统不必追求一次性返回完整长文本，而应优先保证SSE连接稳定和增量消息持续返回；对于统计页和错题趋势页，则应通过分页、聚合查询和缓存策略控制查询开销。")

    add_heading(doc, "4 系统功能设计")
    add_heading(doc, "4.1 系统总体设计")
    add_heading(doc, "4.1.1 总体架构设计")
    add_body(doc, "本系统采用前后端分离架构。前端使用Vue3实现页面交互，后端由SpringBoot提供RESTful接口与流式AI接口，核心业务数据存放于MySQL中，练习草稿与短时状态存放于Redis中，图片等对象资源通过MinIO管理，AI分析与资源推荐能力由千问兼容接口和Dify工作流提供。")
    add_body(doc, "从部署方式看，前端和后端运行在Windows宿主机上，Redis、MinIO和Dify运行在Docker Desktop中。这种架构既符合当前项目的真实运行方式，也便于本地联调和课程展示。")
    add_figure_placeholder(doc, "图4-1 系统总体架构图", "此处插入系统总体架构图，可体现前端、后端、MySQL、Redis、MinIO与AI服务之间的调用关系。")
    add_body(doc, "总体架构设计的核心思想是将页面交互、业务处理、数据存储和智能分析分层组织，使系统在满足当前项目需要的同时，保留后续继续扩展的空间。")

    add_heading(doc, "4.1.2 总体功能设计")
    add_body(doc, "在总体功能层面，系统围绕学生学习闭环与后台内容治理两条主线展开。学生端覆盖学科加入、目录与考点浏览、练习、错题、考试和学情分析；后台管理端覆盖学科、目录、考点、题目、试卷和记录维护；AI能力则贯穿错题复盘、题目问答和学习建议等多个环节。")
    add_figure_placeholder(doc, "图4-2 系统功能模块图", "此处插入系统功能模块图，可展示用户与权限、学科与目录、练习闭环、AI分析、模拟考试和学情统计等模块结构。")

    add_heading(doc, "4.2 系统模块设计")
    add_heading(doc, "4.2.1 用户与权限模块")
    add_body(doc, "用户与权限模块负责注册、登录、获取个人信息和角色鉴权，是整个系统的入口模块。该模块的设计目标是保证普通用户能够顺利进入学习流程，同时通过权限控制明确后台管理接口的访问边界。")
    add_heading(doc, "4.2.2 学科与目录模块")
    add_body(doc, "学科与目录模块负责把题目资源组织成学生可理解、管理员可维护的结构。系统先通过学科建立一级分类，再通过教材目录与考点标签构成章节化和知识点化的双重组织方式，为后续练习、考试和统计提供基础。")
    add_body(doc, "从后台治理角度看，管理员通过AdminSubjectController、AdminDirectoryController、AdminExamTagController和AdminDirectoryTagRelationController完成学科、目录、考点及其关系的维护，使学生端看到的资源范围、章节结构和知识点映射保持一致。")
    add_heading(doc, "4.2.3 练习闭环模块")
    add_body(doc, "练习闭环模块是系统的核心模块，负责开始练习、拉取题单、保存草稿、批量交卷、生成练习记录、沉淀错题和更新统计结果。该模块不仅要完成一次练习过程，还要保证交卷后错题本、统计页和掌握度数据能够同步变化。")
    add_heading(doc, "4.2.4 AI智能分析模块")
    add_body(doc, "AI智能分析模块围绕题目和错题提供增强反馈。系统在错题场景中保存分析快照和追问会话，在练习题场景中提供题目助教问答，在薄弱点分析场景中可进一步调用Dify工作流检索外部资源。")
    add_heading(doc, "4.2.5 模拟考试模块")
    add_body(doc, "模拟考试模块负责试卷配置、配题、学生作答、交卷、成绩汇总和历史记录回看。该模块与练习模块共享题库基础，但在业务上更强调试卷模板、分值管理和记录归档。")
    add_body(doc, "在管理员端，AdminExamController负责试卷列表、试卷新增修改、启停控制、可选题目筛选、配题和考试记录查询等接口，使后台能够持续为学生端提供可用的模拟考试资源。")
    add_heading(doc, "4.2.6 学情统计模块")
    add_body(doc, "学情统计模块负责按用户、按学科、按日期和按知识点输出学习结果，包括累计答题量、正确率、错题趋势、掌握度等级和薄弱点推荐等内容，是业务记录与AI建议之间的重要桥梁。")

    add_heading(doc, "4.3 数据库设计")
    add_heading(doc, "4.3.1 数据库E-R图")
    add_body(doc, "概念结构设计是数据库设计的重要起点，其核心目标是从业务角度梳理系统中的数据实体、实体间联系以及关键属性组织形式。结合当前项目SQL脚本与业务流程，本文采用E-R模型对用户、学科、教材目录、考点标签、题目、练习、考试和AI分析等核心对象进行建模。")
    add_figure_placeholder(doc, "图4-3 总体E-R图", "此处插入系统总体E-R图，可体现用户、学科、目录、考点、题目、练习、考试与AI分析之间的关系。")
    add_body(doc, "基于系统总体实体关系图，可进一步对主要实体的属性结构进行拆解说明。为了与模板中的数据库设计写法保持一致，本文继续对用户、学科、目录、考点、题目、练习、错题、试卷与考试记录等核心实体进行属性级展示，具体如图4-4至图4-14所示。")
    for spec in ER_ENTITY_SPECS:
        add_body(doc, spec["intro"])
        add_figure_placeholder(
            doc,
            spec["caption"],
            f"此处插入{spec['entity_name']}实体属性图，展示该实体的关键字段结构。",
            width_cm=spec["width_cm"],
        )

    add_heading(doc, "4.3.2 数据库表设计")
    add_body(doc, "根据当前SQL脚本，系统已经定义多张核心业务表。为了突出对学习闭环最关键的数据结构，本文对其中主要业务表进行摘要说明。")
    add_caption(doc, "表4-1 核心数据表设计表")
    add_table(
        doc,
        [
            ["表名", "作用", "关键字段"],
            ["user", "保存系统用户基础信息并区分管理员身份", "phone、password、is_admin、status"],
            ["subject", "保存学科基础信息", "id、name、status"],
            ["textbook_directory", "保存教材目录树结构", "subject_id、parent_id、name、sort"],
            ["exam_tag", "保存考点标签信息", "subject_id、name、tag"],
            ["question", "保存题目主体数据", "subject_id、directory_id、type、difficulty、answer"],
            ["practice_session", "保存一次练习会话的总体信息", "user_id、subject_id、practice_mode、status"],
            ["exercise_record", "保存每道练习题的作答明细", "session_ref_id、question_id、user_answer、is_correct"],
            ["wrong_question", "保存用户错题本数据", "user_id、question_id、wrong_count、master_status"],
            ["exam_record", "保存一次考试的总记录", "user_id、subject_id、score、grading_status"],
            ["wrong_question_ai_analysis", "保存错题AI分析快照", "wrong_question_id、result_json、summary、total_tokens"],
        ],
    )

    add_heading(doc, "5 系统功能实现")
    add_heading(doc, "5.1 系统功能实现")
    add_body(doc, "为避免学生端与管理员端内容混杂，本章按照学生端学习功能、管理员端治理功能和公共支撑能力三个角度分别说明系统的实现方式。")
    add_heading(doc, "5.1.1 学生端学科与首页功能实现")
    add_body(doc, "在学科与首页功能方面，系统通过UserSubjectController提供学科目录查询、我的学科列表、加入学科以及学科目录和考点查询接口。学生登录后，可先在学科页查看可加入的学科，再进入对应学科浏览目录树与考点标签，首页则围绕最近学习入口、错题入口和考试入口组织页面跳转。")
    add_body(doc, "该实现使学生端首页不再只是简单导航页，而成为学习主链路的分发入口。只要用户加入学科，就可以围绕学科、目录和考点进入具体练习场景。")
    add_figure_placeholder(
        doc,
        "图5-1 学生端学科与首页功能界面图",
        "此处插入学生端首页或学科页截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/subject/controller/UserSubjectController.java 的 listMySubjects 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-2 学生端学科与首页核心代码图",
        "此处插入首页或学科功能相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/subject/service/impl/UserSubjectServiceImpl.java 的 listMySubjects 方法。",
    )

    add_heading(doc, "5.1.2 学生端练习、错题与AI助教功能实现")
    add_body(doc, "练习闭环功能主要由PracticeController和PracticeServiceImpl承担实现。系统支持开始练习、获取题单、保存草稿、提交整套练习、查询历史会话、查看错题、更新错题掌握状态以及查看错题趋势等能力，学生端的日常刷题过程由此形成稳定主链路。")
    add_body(doc, "在交卷实现中，系统不仅返回答题结果，还会同步写入exercise_record、wrong_question、user_stats、user_subject_stats、daily_stats和user_knowledge_mastery等多类数据。进一步地，WrongQuestionAiController和WrongQuestionAiServiceImpl会读取错题、标准答案、解析和最近作答记录生成AI分析，并支持围绕当前错题继续追问，使学生端形成“练习-错题-分析-再练”的闭环。")
    add_figure_placeholder(
        doc,
        "图5-3 学生端练习、错题与AI助教功能界面图",
        "此处插入练习作答页、错题页或AI助教区域截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/practice/controller/PracticeController.java 的 startPractice 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-4 学生端练习、错题与AI助教核心代码图",
        "此处插入AI分析接口、流式会话或服务实现相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/ai/service/impl/WrongQuestionAiServiceImpl.java 的 analyzeWrongQuestion 方法。",
    )

    add_heading(doc, "5.1.3 学生端模拟考试功能实现")
    add_body(doc, "模拟考试功能由管理员端和学生端共同完成。管理员通过AdminExamController维护试卷、配题和可选题目列表，学生端则通过考试接口选择试卷、开始考试、提交答案并回看考试记录。")
    add_body(doc, "对于主观题场景，系统还保留了AI辅助评分任务与评分状态字段，为后续更复杂的评分流程预留了实现基础。")
    add_figure_placeholder(
        doc,
        "图5-5 学生端模拟考试功能界面图",
        "此处插入考试中心、考试作答页或考试记录页截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/exam/controller/ExamController.java 的 startExam 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-6 学生端模拟考试核心代码图",
        "此处插入试卷配置、考试提交或考试记录相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/exam/service/impl/ExamServiceImpl.java 的 submitExam 方法。",
    )

    add_heading(doc, "5.1.4 学生端学情分析与薄弱点推荐实现")
    add_body(doc, "学情分析功能围绕PracticeController中的统计接口和AI分析控制器展开。系统能够按照用户、学科和时间范围输出累计答题量、正确率、错题趋势和薄弱知识点列表，并结合knowledge_mastery等数据判断某个考点的掌握等级。")
    add_body(doc, "在此基础上，系统还能进一步为薄弱点提供AI解释与视频检索结果，使统计结果从单纯的数据展示转化为可执行的学习建议。")
    add_figure_placeholder(
        doc,
        "图5-7 学生端学情分析与薄弱点推荐界面图",
        "此处插入学情统计页、薄弱点推荐页或图表截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/practice/controller/PracticeController.java 的 getPracticeStats 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-8 学生端学情分析与薄弱点推荐核心代码图",
        "此处插入统计聚合、掌握度计算或推荐相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/practice/service/impl/PracticeServiceImpl.java 的 getPracticeStats 方法。",
    )

    add_heading(doc, "5.1.5 管理员端学科、目录与考点管理实现")
    add_body(doc, "管理员端基础资源治理功能由AdminSubjectController、AdminDirectoryController、AdminExamTagController和AdminDirectoryTagRelationController共同实现。管理员登录后台后，可以维护学科基本信息、目录树结构、考点标签以及目录与考点之间的关联关系，从而决定学生端能够看到什么学科、哪些章节以及每个章节下对应哪些知识点。")
    add_body(doc, "这一部分并不是简单的数据录入页面，而是整个题库平台的内容组织层。尤其是目录考点关联管理，它把教材章节、知识点标签、关联类型、重要程度和考试频率等信息统一在后台治理，使学生端练习、考试和统计都建立在同一套资源结构之上。")
    add_figure_placeholder(
        doc,
        "图5-9 管理员端学科、目录与考点管理界面图",
        "此处插入后台学科管理、目录管理或考点管理页面截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/subject/controller/AdminSubjectController.java 的 list 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-10 管理员端学科、目录与考点管理核心代码图",
        "此处插入后台目录树、考点维护或目录考点关联相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/directory/controller/AdminDirectoryTagRelationController.java 的 list 方法。",
    )

    add_heading(doc, "5.1.6 管理员端题库、试卷与考试记录管理实现")
    add_body(doc, "管理员端题库与考试治理功能由AdminQuestionController和AdminExamController承担实现。题库管理支持按学科、目录、难度、题型和关键字筛选题目，并完成题目的新增、编辑、删除与详情查看；试卷管理则支持试卷列表维护、试卷启停、按目录和考点筛选可用题目并完成配题。")
    add_body(doc, "除了题目与试卷本身，后台还承担考试记录管理职责。管理员可以按学科、用户、关键字和时间范围检索考试记录，并查看某次考试的总览信息与题目明细。这样一来，学生端的模拟考试不再只是前台功能，而是建立在后台完整的试卷生产、配题和记录治理能力之上。")
    add_figure_placeholder(
        doc,
        "图5-11 管理员端题库、试卷与考试记录管理界面图",
        "此处插入后台题库管理、试卷管理或考试记录管理页面截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/question/controller/AdminQuestionController.java 的 list 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-12 管理员端题库、试卷与考试记录管理核心代码图",
        "此处插入后台试卷配题、考试记录查询或权限校验相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/modules/exam/service/impl/ExamServiceImpl.java 的 listPaperAvailableQuestions 方法。",
    )

    add_heading(doc, "5.1.7 文件上传与资源访问实现")
    add_body(doc, "文件上传与资源访问由FileController、MinioConfig和MinioService共同实现。系统通过配置好的MinIO endpoint和bucket完成题图、作答图片等资源的上传、预览和访问地址生成，使图片文件不必直接存放在数据库中。")
    add_body(doc, "这一实现不仅改善了文件类资源的管理方式，也为后续主观题作答、题目图片展示和AI图像评分等能力提供了统一入口。")
    add_figure_placeholder(
        doc,
        "图5-13 文件上传与资源访问功能界面图",
        "此处插入上传结果、图片访问或对象资源页面截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/common/minio/controller/FileController.java 的 upload 方法。",
    )
    add_figure_placeholder(
        doc,
        "图5-14 文件上传与资源访问核心代码图",
        "此处插入MinIO配置或文件上传相关核心代码截图。",
        "后端代码位置：aishua-master/src/main/java/zysy/iflytek/aishua/common/minio/service/MinioService.java 的 upload 方法。",
    )

    add_heading(doc, "6 系统部署上线")
    add_heading(doc, "6.1 系统部署规划")
    add_body(doc, "当前项目采用本地混合部署方式：前端和后端运行在Windows宿主机上，Redis、MinIO和Dify运行在Docker Desktop中，数据库运行在本机MySQL服务。由于本次部署主要面向本地开发与课程展示，因此统一使用localhost访问，不单独配置公网域名。")
    add_caption(doc, "表6-1 系统部署规划表")
    add_table(
        doc,
        [
            ["组件", "部署位置", "地址/端口", "说明"],
            ["前端项目", "Windows宿主机", "http://localhost:8084", "提供学生端与后台管理端页面。"],
            ["后端服务", "Windows宿主机", "http://localhost:8080", "提供统一业务接口与鉴权能力。"],
            ["MySQL", "Windows宿主机", "localhost:3306", "存储题库、练习、考试及分析数据。"],
            ["Redis", "Docker Desktop容器", "localhost:6379", "缓存练习草稿与短时状态。"],
            ["MinIO", "Docker Desktop容器", "http://localhost:9000", "存储题图、作答图片等对象文件。"],
            ["Dify", "Docker Desktop容器", "http://localhost", "提供工作流式AI能力。"],
        ],
        column_widths=[2.2, 3.4, 4.0, 5.4],
        font_size=10.5,
        header_font_size=11,
        line_spacing=14,
        column_alignments=["center", "center", "center", "left"],
    )

    add_heading(doc, "6.2 系统部署操作步骤")
    add_body(doc, "系统部署前，应确认本机已安装JDK 17、Node.js、MySQL 8.0和Docker Desktop，并确保Redis、MinIO和Dify等基础服务能够正常启动。部署操作步骤如下。")
    add_body(doc, "第一步，导入数据库脚本并创建aishua数据库。若数据库尚未初始化，可在宿主机命令行中执行如下命令。")
    add_code_block(
        doc,
        [
            "mysql -uroot -proot -e \"create database if not exists aishua default character set utf8mb4;\"",
            f"mysql -uroot -proot aishua < \"{ROOT_DIR / 'aishua-master' / 'src' / 'main' / 'resources' / 'db' / 'aishua.sql'}\"",
        ],
    )
    add_body(doc, "第二步，检查后端配置文件application.yml，确认MySQL、Redis、MinIO和Dify等地址与本机环境一致，当前项目默认使用localhost相关地址。")
    add_body(doc, "第三步，确认Docker Desktop中的Redis、MinIO和Dify容器已经启动。由于容器名称可能因个人环境不同而变化，建议先查看容器状态，再按实际名称启动。")
    add_code_block(
        doc,
        [
            "docker ps",
            "docker start <redis容器名>",
            "docker start <minio容器名>",
            "docker start <dify容器名>",
        ],
    )
    add_body(doc, "第四步，在后端项目目录启动SpringBoot服务。")
    add_code_block(doc, [f"cd /d \"{ROOT_DIR / 'aishua-master'}\"", "mvn spring-boot:run"])
    add_body(doc, "第五步，在前端项目目录安装依赖并启动开发服务。")
    add_code_block(doc, [f"cd /d \"{ROOT_DIR / 'aishua-web'}\"", "npm install", "npm run serve"])
    add_body(doc, "第六步，当前后端都启动完成后，通过浏览器访问前端地址，检查登录、学科、练习、考试和AI分析等功能是否正常联通。若需要打包部署，也可进一步执行mvn clean package和npm run build。")

    add_heading(doc, "6.3 部署验证")
    add_body(doc, "系统上线验证应覆盖不同角色和不同功能链路。由于当前数据库脚本中未内置固定测试账号，因此普通用户应通过注册流程自行创建账号；若需要验证管理员角色，可将已注册用户的user表记录中的is_admin字段更新为1后重新登录。")
    add_code_block(doc, ["update user", "set is_admin = 1", "where phone = '已注册手机号';"])
    add_caption(doc, "表6-2 上线验证表")
    add_table(
        doc,
        [
            ["验证对象", "访问入口", "验证内容"],
            ["游客或未登录用户", "http://localhost:8084", "确认首页、注册页和登录页可正常访问。"],
            ["普通学生用户", "登录后进入学科页", "完成加入学科、开始练习、保存草稿、提交作答与查看结果。"],
            ["管理员用户", "登录后进入后台管理端", "验证学科、目录、题库、试卷和考试记录管理功能。"],
            ["AI分析功能", "错题页或练习助教入口", "验证AI分析结果、流式消息和会话历史返回正常。"],
            ["文件上传功能", "带图片上传的作答场景", "验证文件可写入MinIO并返回访问地址。"],
            ["基础服务连通性", "8080 / 3306 / 6379 / 9000", "确认前后端与MySQL、Redis、MinIO、Dify链路正常。"],
        ],
        column_widths=[3.2, 4.0, 8.0],
        font_size=10.5,
        header_font_size=11,
        line_spacing=14,
        column_alignments=["center", "center", "left"],
    )

    add_heading(doc, "7 总结与展望")
    add_heading(doc, "7.1 总结")
    add_body(doc, "本文围绕AI刷题系统的设计与实现，完成了从问题提出、需求分析、系统设计到功能实现和部署整理的完整工作。结合项目当前代码和数据库结构，本文不仅说明了系统具备哪些功能，也重点分析了这些功能如何围绕学习闭环形成协同。")
    add_body(doc, "本文完成的主要工作可概括为以下几个方面：")
    add_body(doc, "1. 完成了项目背景与业务问题分析，明确了传统题库系统在资源组织、错题沉淀和智能反馈方面存在的不足。")
    add_body(doc, "2. 完成了系统需求分析，从业务需求、功能需求、数据需求和非功能需求几个层面明确了系统边界。")
    add_body(doc, "3. 完成了系统总体设计与数据库设计，梳理了用户、学科、目录、考点、题目、练习、考试和AI分析等核心实体关系。")
    add_body(doc, "4. 结合当前项目代码，整理了学生端学习闭环与管理员端内容治理两条主线，具体涵盖学科与目录管理、题库与试卷管理、练习闭环、错题AI分析、模拟考试、学情统计和文件上传等模块的实现方式。")
    add_body(doc, "5. 根据项目实际运行方式，给出了Windows宿主机加Docker Desktop混合环境下的部署规划、操作步骤和验证方法。")
    add_body(doc, "6. 形成了一份能够对应模板红字要求的项目设计文档，为后续继续完善正式论文和课程提交材料提供了基础。")
    add_body(doc, "通过本次系统梳理与文档撰写，本人对题库平台的业务建模、学习数据闭环设计、AI能力接入边界以及本地部署联调流程有了更深入的认识，也进一步体会到需求分析与系统结构设计对于项目开发质量的重要作用。")

    add_heading(doc, "7.2 进一步的工作")
    add_body(doc, "虽然本文已经较完整地描述了AI刷题系统的设计与实现，但系统仍存在进一步完善空间。第一，AI分析能力虽然已经能够结合错题上下文输出结果，但在分析稳定性、质量评估和模型成本控制方面仍需继续优化。")
    add_body(doc, "第二，当前题库运营仍以人工维护为主，后续可进一步完善题目批量导入、题目审核、目录与考点关系治理等能力，以提升后台内容维护效率。")
    add_body(doc, "第三，当前部署方案主要面向本地开发与课程展示，若要进一步走向更正式的演示或生产环境，还需补充HTTPS接入、环境变量化配置、日志审计和持续集成等能力。")
    add_body(doc, "第四，当前学情统计与薄弱点推荐已具备基础能力，未来还可继续扩展个性化学习路径推荐、AI补题、长期学习计划生成和资源效果追踪等功能，使系统逐步演化为更完整的智能学习助手平台。")

    add_heading(doc, "附录")
    add_body(doc, "附录部分用于补充正文中未展开的关键接口、重要参数和核心数据表摘要，以便读者快速理解系统的真实落地结构。")
    add_caption(doc, "表A-1 核心接口摘要")
    add_table(
        doc,
        [
            ["模块", "接口", "说明"],
            ["用户认证", "/api/user/login、/api/user/register", "完成普通用户登录与注册，是系统公开接口。"],
            ["学科中心", "/api/subjects、/api/user/subjects/{subjectId}/join", "完成学科展示、加入和我的学科管理。"],
            ["后台学科管理", "/api/admin/subjects", "完成学科列表查询、新增、编辑、启停和删除。"],
            ["后台目录与考点管理", "/api/admin/directories、/api/admin/tags、/api/admin/directory-tag-relations", "完成目录树、考点标签和目录考点关系维护。"],
            ["后台题库管理", "/api/admin/questions", "完成题目筛选、详情查看、新增、编辑和删除。"],
            ["练习模块", "/api/practice/start、/api/practice/{sessionId}/submit-all", "完成开始练习、交卷和结果汇总。"],
            ["草稿缓存", "/api/practice/{sessionId}/draft", "完成练习草稿读取与保存。"],
            ["错题模块", "/api/practice/wrong-questions", "完成错题查询、掌握状态更新和趋势统计。"],
            ["错题AI分析", "/api/practice/wrong-questions/{wrongQuestionId}/ai-analysis", "完成错题分析结果生成。"],
            ["练习题助教", "/api/practice/{sessionId}/questions/{questionId}/assistant/...", "完成围绕单题的AI追问会话。"],
            ["考试模块", "/api/exam/...、/api/admin/exams/...", "完成试卷配置、考试作答和记录查看。"],
            ["后台试卷与记录管理", "/api/admin/exams/papers、/api/admin/exams/records", "完成试卷维护、配题和考试记录查询。"],
            ["文件上传", "/api/files", "完成对象文件上传与资源访问地址生成。"],
        ],
    )
    add_caption(doc, "表A-2 关键参数摘要")
    add_table(
        doc,
        [
            ["参数项", "当前值", "作用"],
            ["spring.datasource.url", "jdbc:mysql://localhost:3306/aishua", "指定业务数据库连接地址。"],
            ["spring.redis.host / port", "localhost / 6379", "指定Redis缓存地址。"],
            ["minio.endpoint", "http://localhost:9000", "指定MinIO对象存储访问地址。"],
            ["minio.bucket-name", "subject", "指定对象存储桶名称。"],
            ["dify.base-url", "http://localhost", "指定Dify工作流服务地址。"],
            ["app.security.jwt.expiration-ms", "86400000", "控制JWT令牌有效期。"],
            ["practice.draft.redis-key-prefix", "practice:draft:v2:", "控制练习草稿缓存Key前缀。"],
        ],
    )
    add_caption(doc, "表A-3 核心数据表摘要")
    add_table(
        doc,
        [
            ["表名", "业务含义", "备注"],
            ["user", "系统用户表", "保存手机号、密码、昵称和管理员标识。"],
            ["subject", "学科表", "保存系统中的学科基础信息。"],
            ["textbook_directory", "教材目录表", "保存目录树结构，支撑章节化浏览与筛选。"],
            ["question", "题目表", "保存客观题与主观题主体信息。"],
            ["practice_session", "练习会话表", "保存一次练习的总体元数据。"],
            ["exercise_record", "练习记录表", "保存每道题的作答结果与判题信息。"],
            ["wrong_question", "错题本表", "保存用户错题、错误次数和掌握状态。"],
            ["exam_record", "考试记录表", "保存考试成绩、评分状态与考试过程数据。"],
            ["wrong_question_ai_analysis", "错题AI分析表", "保存结构化分析结果与模型调用信息。"],
        ],
    )
    add_body(doc, "关键源码位置可进一步参考以下文件：UserSubjectController.java、AdminSubjectController.java、AdminDirectoryController.java、AdminDirectoryTagRelationController.java、AdminQuestionController.java、PracticeController.java、PracticeServiceImpl.java、WrongQuestionAiController.java、WrongQuestionAiServiceImpl.java、AdminExamController.java、AdminAccess.java、AuthInterceptor.java、MinioConfig.java以及application.yml。")


def validate_report(output_path: Path) -> None:
    doc = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    lowered = text.lower()
    forbidden = ["<ydh", "智慧公益助盲", "视障求助", "uni-app", "rabbitmq", "文字化"]
    remains = [item for item in forbidden if item.lower() in lowered]
    if remains:
        raise ValueError(f"文档仍包含不应残留内容：{', '.join(remains)}")
    required = ["1.3.2 主要工作", "3.2.2 业务分析", "6.3 部署验证", "附录"]
    missing = [item for item in required if item not in text]
    if missing:
        raise ValueError(f"文档缺少关键章节：{', '.join(missing)}")


def main() -> None:
    template_path = discover_template()
    doc = Document(template_path)
    patch_front_matter(doc)
    rebuild_body(doc)
    doc.save(OUTPUT_PATH)
    validate_report(OUTPUT_PATH)
    print(f"已生成文档：{OUTPUT_PATH}")


if __name__ == "__main__":
    main()
